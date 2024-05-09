from django.shortcuts import render, HttpResponseRedirect, HttpResponse
from dispenda_surat.forms import Pegawai_Form, Jabatan_Form, SuratTugas_Form,Pegawai_Update, Struktur_Form
from dispenda_surat.forms import SuratTugas_Update, Struktur_Update, User_Form, Konfigurasi_Form, SuratTugas_Form
from dispenda_surat.forms import DasarSuratTugas_Form, Evaluasi_Form, Evaluasi_Form_Update
from dispenda_surat.models import Pegawai,Jabatan, Struktur,  Konfigurasi, SuratTugas,DasarSuratTugas
from dispenda_surat.models import ST_Dasar, ST_Peserta, LaporanEval
from django.contrib.auth import authenticate
from django.contrib.auth import login,logout
# from django.contrib.auth.models import User
from django.db.models import Q
from django.conf import settings
import os
import random
import docx
from docx.shared import Pt,Cm
from django.conf import settings


# Create your views here.
def Login(request):
    print('login')
    if(request.user.is_authenticated):
        print('belum login')
        return render(request,'dashboard.html')
    else:
        print('lagi mau login')
        if request.method == "POST":
            print('ini submit')
            x='human' in request.POST
            if x:
                print('ini adalah manusia')
                username=request.POST['username']
                password=request.POST['password']
                user = authenticate(username=username,password=password)
                if(user is not None):
                    print('login berhasil')
                    login(request,user)
                    return HttpResponseRedirect('/')
        return render(request,'login.html')

def Logout(request):
    logout(request)
    return HttpResponseRedirect('/')

def Pegawai_create(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    if request.method == 'POST':
        myform = Pegawai_Form(request.POST)
        if myform.is_valid():
            myform.save()
    
    myform = Pegawai_Form()
    context = {
        'myform': myform,
    }
    return render(request,'pegawai/create.html', context)

def Pegawai_update(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    myform = SuratTugas_Form()
    print("test")
    return render(request,'pegawai/update.html',{'myform':myform})

def Pegawai_list(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    data = Pegawai.objects.all()
    context = {
        'data':data
    }
    return render(request,'pegawai/list.html',context)

def Pegawai_delete(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    return HttpResponseRedirect('/pegawai/list/')

def Jabatan_create(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    if request.method == 'POST':
        myform = Jabatan_Form(request.POST)
        if myform.is_valid():
            myform.save()
    
    myform = Jabatan_Form()
    context = {
        'myform': myform
    }
    return render(request,'jabatan/create.html', context)

def Jabatan_update(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    return render(request,'jabatan/update.html')

def Jabatan_list(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    data = Jabatan.objects.all()
    context = {
        'data':data
    }
    return render(request,'jabatan/list.html',context)

def Jabatan_delete(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    return HttpResponseRedirect('/jabatan/list/')

def Struktur_create(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    if request.method == 'POST':
        myform = Struktur_Form(request.POST)
        if myform.is_valid():
            myform.save()
    
    myform = Struktur_Form()
    context = {
        'myform': myform
    }
    return render(request,'struktur/create.html', context)

def Struktur_update(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    
    return render(request,'struktur/update.html')

def Struktur_list(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    data = Struktur.objects.all()
    context = {
        'data':data
    }
    return render(request,'struktur/list.html',context)

def Struktur_delete(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    return HttpResponseRedirect('/struktur/list/')


def User_create(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    if request.method == 'POST':
        myform = User_Form(request.POST)
        if myform.is_valid():
            myform.save()
            user = RealUser.objects.get(username=request.POST.get('username'))
            user.set_password(request.POST.get('password'))
            user.save()
        else:
            print(myform)
    
    myform = User_Form()
    context = {
        'myform': myform
    }
    return render(request,'user/create.html', context)

def User_update(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    
    return render(request,'struktur/update.html')

def User_list(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    data = RealUser.objects.all()
    context = {
        'data':data
    }
    return render(request,'user/list.html',context)

def User_delete(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    return HttpResponseRedirect('/struktur/list/')

def Konfigurasi_create(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    if(Konfigurasi.objects.all().count()>0):
        return HttpResponseRedirect('/konfigurasi/list/')
    if request.method == 'POST':
        myform = Konfigurasi_Form(request.POST)
        if myform.is_valid():
            myform.save()
            return HttpResponseRedirect('/konfigurasi/list/')
    
    myform = Konfigurasi_Form()
    context = {
        'myform': myform
    }
    return render(request,'konfigurasi/create.html', context)

def Konfigurasi_list(request):
    if(not request.user.is_superuser):
        return HttpResponseRedirect('/')
    if(Konfigurasi.objects.all().count()==0):
        return HttpResponseRedirect('/konfigurasi/create/')
    
    data = Konfigurasi.objects.all()[0]
    context = {
        'data': data
    }
    return render(request,'konfigurasi/list.html', context)

def SuratTugas_create(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    
    isGet = False
    try:
        nomorsurat = request.GET['surat']
        isGet=True
    except:
        pass

    if(isGet):
        surattugas = SuratTugas.objects.all().get(Nomor_Surat=nomorsurat)
        
        peserta = ST_Peserta.objects.filter(Nomor_Surat=surattugas)
        dasar = ST_Dasar.objects.filter(Nomor_Surat=surattugas)
        

        list_peserta = Pegawai.objects.all().filter(~Q(Nama=surattugas.Ketua_Tim) & ~Q(Nama=surattugas.Kepala_Bapeda))
        list_dasar = DasarSuratTugas.objects.all()
        
        context = {
                'surattugas':surattugas,
                'peserta': peserta,
                'dasar':dasar,
                'list_peserta':list_peserta,
                'list_dasar':list_dasar,
        }
        return render(request,'surattugas/create_detail.html',context)
    
    if(request.method == 'POST'):
        form = SuratTugas_Form(request.POST)
        if(form.is_valid()):
            form.save()
            Kepala_Bapedanya = Konfigurasi.objects.all()[0].Kepala_Bapeda.Nama
            SuratTugas.objects.filter(Nomor_Surat = request.POST['Nomor_Surat']).update(Kepala_Bapeda=Kepala_Bapedanya, ID_NomorSurat=random.random()*1000000)
            return HttpResponseRedirect('/surattugas/create/?surat=' + request.POST['Nomor_Surat'])
    form = SuratTugas_Form()
    context = {
        'form':form
    }
    return render(request,'surattugas/create.html', context)

def SuratTugas_list(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    data = SuratTugas.objects.all()
    context = {
        'data':data
    }
    return render(request,'surattugas/list.html', context)

def DasarTugas_create(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    if(request.method == 'POST'):
        form = DasarSuratTugas_Form(request.POST)
        if(form.is_valid()):
            form.save()
            return HttpResponseRedirect('/dasartugas/list/')
    form = DasarSuratTugas_Form()
    context = {
        'form':form
    }
    return render(request,'dasarnya/create.html', context)

def DasarTugas_list(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    data = DasarSuratTugas.objects.all()
    context = {
        'data':data
    }
    return render(request,'dasarnya/list.html', context)

def Anggota_Tambah(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    try:
        nomorsurat = request.GET['surat']
    except:
        return HttpResponseRedirect('/surattugas/list/')
    try:
        if SuratTugas.objects.get(Nomor_Surat=nomorsurat).isDone==True:
            return HttpResponseRedirect('/surattugas/list/')
        st_peserta = ST_Peserta()
        #.objects.create(Dasar=DasarSuratTugas.objects.get(dasar=request.POST['dasarnya']),Surat_Tugas=SuratTugas.objects.get(Nomor_Surat=nomorsurat))
        st_peserta.Nomor_Surat = SuratTugas.objects.get(Nomor_Surat=nomorsurat)
        st_peserta.Peserta = Pegawai.objects.get(NIP=request.POST['pesertanya'])
        st_peserta.save()
    except:
        pass
    return HttpResponseRedirect('/surattugas/create/?surat=' + nomorsurat)

def Dasar_Tambah(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    try:
        nomorsurat = request.GET['surat']
    except:
        return HttpResponseRedirect('/surattugas/list/')
    
    try:
        if SuratTugas.objects.get(Nomor_Surat=nomorsurat).isDone==True:
            return HttpResponseRedirect('/surattugas/list/')
        st_dasar = ST_Dasar()
        #.objects.create(Dasar=DasarSuratTugas.objects.get(dasar=request.POST['dasarnya']),Surat_Tugas=SuratTugas.objects.get(Nomor_Surat=nomorsurat))
        st_dasar.Nomor_Surat = SuratTugas.objects.get(Nomor_Surat=nomorsurat)
        st_dasar.Dasar = DasarSuratTugas.objects.get(dasar=request.POST['dasarnya'])
        st_dasar.save()
    except:
        pass
    return HttpResponseRedirect('/surattugas/create/?surat=' + nomorsurat)

def Exportkan(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    try:
        nomorsurat = request.GET['nosur']
        surattugas = SuratTugas.objects.get(ID_NomorSurat=nomorsurat)
    except:
        nomorsurat = None

    if nomorsurat is not None:
        document = docx.Document()

        font = document.styles['Normal'].font
        font.name = "Time New Roman"
        font.size = Pt(12)

        sections = document.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.bottom_margin = Cm(2)

        document.add_picture(os.path.join(settings.BASE_DIR,'static/img/kop_surat.png'),width=Cm(17), height=Cm(4))

        paragraph = document.add_paragraph()
        runner = paragraph.add_run("SURAT PERINTAH TUGAS")
        runner.bold=True
        runner.underline=True
        runner = paragraph.add_run("\nNomor: %s"%surattugas.Nomor_Surat)
        paragraph.alignment = 1

        table = document.add_table(rows=1,cols=2)
        row = table.rows[0].cells
        paragraph=row[0].add_paragraph()
        runner = paragraph.add_run("DASAR:")
        runner.bold=True
        row[1].add_paragraph("Keputusan Menteri Dalam Negeri Nomor 16 Tahun 2013 tanggal 23 Januari 2013 tentang Pelaksanaan Perjalanan Dinas;",style="List Number")
        row[1].add_paragraph("Peraturan Daerah Provinsi Jawa Tengah Nomor 12 Tahun 2021 tentang Anggaran Pendapatan dan Belanja Daerah Provinsi Jawa Tengah Tahun Anggaran 2022;",style="List Number")
        row[1].add_paragraph("Peraturan Gubernur Jawa Tengah Nomor 17 Tahun 2013 tentang Perjalanan Dinas Gubernur / Wakil Gubernur, Pimpinan dan Anggota Dewan Perwakilan Rakyat Daerah, Pegawai Negeri Sipil, Calon Pegawai Negeri Sipil Dan Pegawai Non Pegawai Negeri Sipil;",style="List Number")
        row[1].add_paragraph("Peraturan Gubernur Jawa Tengah Nomor 27 Tahun 2020  tentang Standar Harga Satuan Provinsi Jawa;",style="List Number")
        print(surattugas)
        #lookup for dasar tambahan
        dasar = ST_Dasar.objects.all().filter(Nomor_Surat=SuratTugas.objects.get(ID_NomorSurat=nomorsurat))
        if(dasar.count()>0):
            for dasarnya in dasar:
                row[1].add_paragraph(dasarnya.Dasar.dasar + ";",style="List Number")

        for cell in table.columns[0].cells:
            cell.width = Cm(3)
        for cell in table.columns[1].cells:
            cell.width = Cm(14)
        
        paragraph = document.add_paragraph()
        runner = paragraph.add_run("\nMEMERINTAHKAN")
        runner.bold=True
        paragraph.alignment = 1

        table = document.add_table(rows=1,cols=2)
        row = table.rows[0].cells
        
        paragraph = row[0].add_paragraph()
        runner = paragraph.add_run("KEPADA:")
        runner.bold=True
        parnya = row[1].add_paragraph("1.   Nama:\t" + surattugas.Ketua_Tim.Nama)
        runner = parnya.add_run("\n      NIP:\t" + surattugas.Ketua_Tim.NIP)
        runner = parnya.add_run("\n      Jabatan:\t" + surattugas.Ketua_Tim.Jabatan.Nama_Jabatan)


        #lookup for dasar tambahan
        peserta = ST_Peserta.objects.all().filter(Nomor_Surat=SuratTugas.objects.get(ID_NomorSurat=nomorsurat))
        counter = 2
        if(peserta.count()>0):
            for pesertanya in peserta:
                parnya = row[1].add_paragraph(str(counter)+".   Nama:\t" + pesertanya.Peserta.Nama)
                runner = parnya.add_run("\n      NIP:\t" + pesertanya.Peserta.NIP)
                runner = parnya.add_run("\n      Jabatan:\t" + pesertanya.Peserta.Jabatan.Nama_Jabatan)

        for cell in table.columns[0].cells:
            cell.width = Cm(3)
        for cell in table.columns[1].cells:
            cell.width = Cm(14)

        table = document.add_table(rows=1,cols=2)
        row = table.rows[0].cells
        paragraph=row[0].add_paragraph()
        runner = paragraph.add_run("UNTUK:")
        runner.bold=True
        mypar=row[1].add_paragraph("")
        runner = mypar.add_run("1.   Melaksanakan tugas:\t%s\n"%surattugas.Tugas)
        runner = mypar.add_run("      di:\t\t\t\t%s\n"%surattugas.Lokasi_Tugas)
        runner = mypar.add_run("      tanggal:\t\t\t%s\n"%surattugas.Tanggal_Tugas)
        runner = mypar.add_run("2.   Tidak menerima gratifikasi dalam bentuk apapun sesuai ketentuan;\n")
        runner = mypar.add_run("3.   Melapor kepada Pejabat setempat guna pelaksanaan tugas tersebut;\n")
        runner = mypar.add_run("4.   Melaporkan Hasil Pelaksanaan Tugas kepada Pejabat pemberi tugas.")
        print(surattugas)
           
        for cell in table.columns[0].cells:
            cell.width = Cm(3)
        for cell in table.columns[1].cells:
            cell.width = Cm(14)

        paragraph = document.add_paragraph()
        runner = paragraph.add_run("\n\t\t\t\t\t\t\tDitetapkan di:\tSemarang\n")
        runner = paragraph.add_run("\t\t\t\t\t\t\tpada tanggal:\t\t%s\n" %surattugas.Tanggal_Surat)
        runner = paragraph.add_run("\t\t\t\t\t\t\t___________________________________________________\n")
        PLT = Konfigurasi.objects.all()[0].isPLT
        if(PLT):
            PLT="Plt. "
        else:
            PLT=""
        runner = paragraph.add_run(PLT + "\t\t\t\t\t\t\tKEPADA BADAN PENGELOLA PENDAPATAN\n")
        runner = paragraph.add_run("\t\t\t\t\t\t\tDAERAH PROVINSI JAWA TENGAH\n")
        runner = paragraph.add_run("\t\t\t\t\t\t\t"+Konfigurasi.objects.all()[0].Kepala_Bapeda.Jabatan.Nama_Jabatan)
        runner = paragraph.add_run("\n\n\n\n\n\n")
        runner = paragraph.add_run("\t\t\t\t\t\t\t")
        runner = paragraph.add_run(Konfigurasi.objects.all()[0].Kepala_Bapeda.Nama)
        runner.underline=True
        runner = paragraph.add_run("\n\t\t\t\t\t\t\t"+Konfigurasi.objects.all()[0].Kepala_Bapeda.Struktur.Nama_Struktur)

        document.save(os.path.join(settings.BASE_DIR,str(nomorsurat) + '.docx'))

        document = open(os.path.join(settings.BASE_DIR,str(nomorsurat) + '.docx'),'rb')
 
        resp = HttpResponse(document,content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = str(nomorsurat) + '.docx'
        resp['Content-Disposition']='attachment;filename=%s'%format(filename)
        return resp
    else:
        return HttpResponseRedirect('/surattugas/list/')
    
def SuratTugas_Done(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    
    try:
        nosurat = request.GET.get('nosurat')

    except:
        return HttpResponseRedirect('/surattugas/list/')
    
    SuratTugas.objects.all().filter(Nomor_Surat=nosurat).update(isDone=True)
    
    return HttpResponseRedirect('/surattugas/list/')

def SuratTugas_Delete(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    try:
        nosurat = request.GET.get('nosurat')

    except:
        return HttpResponseRedirect('/surattugas/list/')
    
    ST_Dasar.objects.all().filter(Nomor_Surat=SuratTugas.objects.get(Nomor_Surat=nosurat)).delete()
    ST_Peserta.objects.all().filter(Nomor_Surat=SuratTugas.objects.get(Nomor_Surat=nosurat)).delete()
    SuratTugas.objects.all().filter(Q(Nomor_Surat=nosurat) & ~Q(isDone=True)).delete()
    
    return HttpResponseRedirect('/surattugas/list/')

def Evaluasi_create(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    try:
        nosur=request.GET.get('surat')
        instance = LaporanEval.objects.get(Nomor_Surat_Eval=nosur)
        
        if request.method == 'POST':
            form = Evaluasi_Form_Update(request.POST,instance=instance)
            if form.is_valid():
                form.save()
            nosur_tugas = instance.Nomor_Surat_Tugas.Nomor_Surat
            ketua_tim = SuratTugas.objects.get(Nomor_Surat=nosur_tugas)
            anggota_tim = ST_Peserta.objects.all().filter(Nomor_Surat=ketua_tim)
            context = {
                'evaluasi':instance,
                'ketua_tim':ketua_tim,
                'anggota_tim':anggota_tim,
            }
            print("berhasil")
            return render(request, 'evaluasi/create_detail.html',context)

        form = Evaluasi_Form_Update(instance=instance)
        context = {
            'form':form
        }
        return render(request,'evaluasi/create.html', context)
    except:
        pass
    if(request.method == 'POST'):
        form = Evaluasi_Form(request.POST)
        if(form.is_valid()):
            form.save()
            instance = LaporanEval.objects.get(Nomor_Surat_Eval=request.POST.get('Nomor_Surat_Eval'))
            nosur_tugas = instance.Nomor_Surat_Tugas.Nomor_Surat
            ketua_tim = SuratTugas.objects.get(Nomor_Surat=nosur_tugas)
            anggota_tim = ST_Peserta.objects.all().filter(Nomor_Surat=ketua_tim)
            context = {
                'evaluasi':instance,
                'ketua_tim':ketua_tim,
                'anggota_tim':anggota_tim,
            }
            print("berhasil")
            return render(request, 'evaluasi/create_detail.html',context)
    form = Evaluasi_Form()
    context = {
        'form':form
    }
    return render(request,'evaluasi/create.html', context)

def Evaluasi_list(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    data = LaporanEval.objects.all()
    context = {
        'data':data
    }
    return render(request,'evaluasi/list.html', context)

def Evaluasi_Delete(request):
    if(not request.user.is_authenticated):
        return HttpResponseRedirect('/')
    try:
        nosurat = request.GET.get('nosurat')

    except:
        return HttpResponseRedirect('/evaluasi/list/')
    
    LaporanEval.objects.all().filter(~Q(isDone=True) & Q(Nomor_Surat_Eval=nosurat)).delete()
    
    return HttpResponseRedirect('/evaluasi/list/')